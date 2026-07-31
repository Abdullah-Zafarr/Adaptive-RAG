import os
import hashlib
import requests
from bs4 import BeautifulSoup
from typing import List, Dict, Any
import pypdf
import docx

# Helper utilities
def get_file_hash(file_bytes: bytes) -> str:
    """Generate SHA256 hash of file content for unique tracking."""
    return hashlib.sha256(file_bytes).hexdigest()

# Custom document object structure (replaces LangChain Document)
class NativeDocument:
    """Native Document container replacing LangChain Document."""
    def __init__(self, page_content: str, metadata: Dict[str, Any] = None):
        self.page_content = page_content
        self.metadata = metadata or {}

    def __repr__(self):
        return f"NativeDocument(len={len(self.page_content)}, metadata={self.metadata})"

# Parses PDF, TXT, DOCX, and Web URL inputs into NativeDocument lists
class DocumentLoaderTool:
    """Native Document Loader Tool for PDF, TXT, DOCX, and Web URLs without LangChain."""

    # Extracts text per page from PDF files (with OCR fallback)
    @staticmethod
    def load_pdf(file_path: str) -> List[NativeDocument]:
        """Load text and metadata from digital and scanned image PDF files (with OCR fallback)."""
        docs = []
        filename = os.path.basename(file_path)
        reader = pypdf.PdfReader(file_path)
        for page_idx, page in enumerate(reader.pages):
            text = page.extract_text() or ""

            # OCR Fallback for scanned/image pages if text is empty
            if not text.strip():
                try:
                    import pdfplumber
                    with pdfplumber.open(file_path) as pdf:
                        if page_idx < len(pdf.pages):
                            p = pdf.pages[page_idx]
                            text = p.extract_text() or ""
                except Exception:
                    pass

            if not text.strip():
                try:
                    import pytesseract
                    from PIL import Image
                    import fitz  # PyMuPDF
                    doc_fitz = fitz.open(file_path)
                    page_fitz = doc_fitz[page_idx]
                    pix = page_fitz.get_pixmap()
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    text = pytesseract.image_to_string(img) or ""
                except Exception:
                    pass

            if text.strip():
                docs.append(NativeDocument(
                    page_content=text,
                    metadata={
                        "source": file_path,
                        "filename": filename,
                        "file_type": "pdf",
                        "page": page_idx + 1
                    }
                ))
        return docs

    # Reads raw text files with UTF-8 encoding
    @staticmethod
    def load_txt(file_path: str) -> List[NativeDocument]:
        """Load text from TXT files."""
        filename = os.path.basename(file_path)
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()
        return [NativeDocument(
            page_content=content,
            metadata={
                "source": file_path,
                "filename": filename,
                "file_type": "txt",
                "page": 1
            }
        )]

    # Reads Microsoft Word documents using python-docx
    @staticmethod
    def load_docx(file_path: str) -> List[NativeDocument]:
        """Load text from DOCX files using python-docx."""
        doc = docx.Document(file_path)
        full_text = [p.text for p in doc.paragraphs if p.text.strip()]
        content = "\n\n".join(full_text)
        filename = os.path.basename(file_path)
        return [NativeDocument(
            page_content=content,
            metadata={
                "source": file_path,
                "filename": filename,
                "file_type": "docx",
                "page": 1
            }
        )]

    # Fetches webpage, strips HTML boilerplate, extracts clean text
    @staticmethod
    def load_url(url: str) -> List[NativeDocument]:
        """Load content from web URL using requests and BeautifulSoup4."""
        headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)"}
        resp = requests.get(url, headers=headers, timeout=10)
        resp.raise_for_status()
        
        soup = BeautifulSoup(resp.text, "html.parser")
        # Remove script and style elements
        for element in soup(["script", "style", "nav", "footer", "header"]):
            element.decompose()
            
        text = soup.get_text(separator="\n")
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        cleaned_text = "\n".join(lines)
        
        return [NativeDocument(
            page_content=cleaned_text,
            metadata={
                "source": url,
                "filename": url,
                "file_type": "url",
                "page": 1
            }
        )]

    # Routes file path to correct parser based on extension
    @classmethod
    def load_file(cls, file_path: str) -> List[NativeDocument]:
        """Auto-detect extension and load file into NativeDocument objects."""
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".pdf":
            return cls.load_pdf(file_path)
        elif ext == ".txt":
            return cls.load_txt(file_path)
        elif ext in [".docx", ".doc"]:
            return cls.load_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")


